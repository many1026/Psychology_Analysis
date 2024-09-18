
# %%
import os
from docx import Document

# Variables para acumular el texto extraído
contenido_sesiones = []
contenido_examenes_medicos = []

def get_all_files(directory):
    # Recorre todas las carpetas de fechas y subcarpetas de pacientes
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx") and not file.startswith("~$"):  # Ignorar archivos temporales de Word
                if not file.startswith("Capturas"):  # Ignorar archivos que comienzan con "Capturas"
                    yield os.path.join(root, file), file

def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        full_text = []

        # Extraer el texto de los párrafos normales
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extraer el texto de las tablas
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                full_text.append(' | '.join(row_data))  # Concatenar el texto de las celdas con separadores

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error procesando archivo {docx_file}: {e}")
        return ""

def process_sesion(file_path):
    # Procesa un archivo de sesión
    doc_text = extract_text_from_docx(file_path)
    contenido_sesiones.append(f"Sesión en {file_path}:\n{doc_text}\n")

def process_examen_medico(file_path):
    # Procesa un archivo de examen médico
    doc_text = extract_text_from_docx(file_path)
    contenido_examenes_medicos.append(f"Examen médico en {file_path}:\n{doc_text}\n")

def classify_and_process_files(directory):
    for file_path, file_name in get_all_files(directory):
        if file_name.startswith("EM"):
            process_examen_medico(file_path)  # Procesar como examen médico
        else:
            process_sesion(file_path)  # Procesar como sesión regular

# Ejecutar procesamiento
directory = "Miriam /"  # Cambia esta ruta a tu carpeta
classify_and_process_files(directory)
#%%
# Imprimir todo el contenido de las sesiones y exámenes médicos al final
print("Contenido de todas las sesiones:\n")
for sesion in contenido_sesiones:
    print(sesion)

print("\nContenido de todos los exámenes médicos:\n")
for examen in contenido_examenes_medicos:
    print(examen)

#%%
import spacy
import re
#%%
# Cargar el modelo de español de SpaCy
nlp = spacy.load("es_core_news_md")

# Función para analizar texto y extraer entidades (como síntomas o diagnósticos)
def extract_key_entities(text):
    doc = nlp(text)
    sintomas = []
    
    # Buscar entidades que podrían ser síntomas o diagnósticos
    for ent in doc.ents:
        if ent.label_ in ["SYMPTOM", "DISEASE", "EMOTION"]:  # Ajusta las etiquetas según el modelo que uses
            sintomas.append(ent.text)
    
    return sintomas

# Función para encontrar técnicas terapéuticas usando regex
def encontrar_tecnicas_terapeuticas(texto):
    tecnicas = re.findall(r"(cognitivo conductual|psicoeducación|terapia familiar|introspección)", texto, re.IGNORECASE)
    return tecnicas
#%%
# Analizar el contenido de las sesiones
print("Análisis de las sesiones:\n")
for sesion in contenido_sesiones:
    print(f"Procesando sesión:\n{sesion[:100]}...")  # Muestra los primeros 100 caracteres de la sesión
    sintomas_encontrados = extract_key_entities(sesion)
    tecnicas_encontradas = encontrar_tecnicas_terapeuticas(sesion)
    
    print(f"Sintomas encontrados: {sintomas_encontrados}")
    print(f"Técnicas terapéuticas encontradas: {tecnicas_encontradas}\n")
#%%
# Analizar el contenido de los exámenes médicos
print("\nAnálisis de los exámenes médicos:\n")
for examen in contenido_examenes_medicos:
    print(f"Procesando examen médico:\n{examen[:100]}...")  # Muestra los primeros 100 caracteres del examen
    sintomas_encontrados = extract_key_entities(examen)
    tecnicas_encontradas = encontrar_tecnicas_terapeuticas(examen)
    
    print(f"Sintomas encontrados: {sintomas_encontrados}")
    print(f"Técnicas terapéuticas encontradas: {tecnicas_encontradas}\n")

# %%
